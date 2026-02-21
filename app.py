import streamlit as st
import datetime
import io
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
import google.generativeai as genai

# --- KONFIGURASI AMAN ---
try:
    GEMINI_API_KEY = st.secrets["GEMINI_API_KEY"]
    genai.configure(api_key=GEMINI_API_KEY)
except Exception:
    st.error("Waduh, API Key belum disetting di Secrets Streamlit!")
    st.stop()

@st.cache_resource
def load_ai_model():
    try:
        available_models = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
        model_name = 'models/gemini-1.5-flash' if 'models/gemini-1.5-flash' in available_models else available_models[0]
        return genai.GenerativeModel(model_name)
    except:
        st.error("Koneksi AI Bermasalah!")
        st.stop()

model = load_ai_model()

st.set_page_config(page_title="Sekretaris PSH Tegal", page_icon="📝")

# --- FUNGSI TANGGAL ---
def get_tanggal_indo():
    bulan_indo = {1: "Januari", 2: "Februari", 3: "Maret", 4: "April", 5: "Mei", 6: "Juni",
                  7: "Juli", 8: "Agustus", 9: "September", 10: "Oktober", 11: "November", 12: "Desember"}
    now = datetime.datetime.now()
    return f"{now.day} {bulan_indo[now.month]} {now.year}"

# --- LOGIKA FORMATTING PRO ---
def format_surat_sekretaris(doc, tag, content):
    """
    Logika cerdas untuk membedakan paragraf narasi dan daftar poin acara.
    """
    for paragraph in doc.paragraphs:
        if tag in paragraph.text:
            # Hapus tag template
            paragraph.text = paragraph.text.replace(tag, "")
            
            lines = content.split('\n')
            for i, line in enumerate(lines):
                clean_line = re.sub(r'[*#_]', '', line).strip()
                if not clean_line: continue
                
                # Buat baris baru di dalam paragraf yang sama atau paragraf baru
                if i > 0:
                    new_p = doc.add_paragraph()
                else:
                    new_p = paragraph
                
                new_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
                # JIKA BARIS ADALAH POIN (Ada titik dua)
                if ":" in clean_line and len(clean_line.split(":")[0]) < 20:
                    label, detail = clean_line.split(":", 1)
                    # Beri tabulasi ke kanan (indent)
                    new_p.paragraph_format.left_indent = Inches(0.5)
                    # Titik dua sejajar di 2.0 inci
                    tab_stops = new_p.paragraph_format.tab_stops
                    tab_stops.add_tab_stop(Inches(2.0), WD_TAB_ALIGNMENT.LEFT)
                    
                    run = new_p.add_run(f"{label.strip()}\t: {detail.strip()}")
                else:
                    # JIKA PARAGRAF NARASI
                    new_p.paragraph_format.left_indent = Inches(0)
                    new_p.paragraph_format.first_line_indent = Inches(0.5) # Baris pertama menjorok
                    run = new_p.add_run(clean_line)
                
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)

# --- UI DASHBOARD ---
st.title("🛡️ Sekretaris Digital PSH Tegal")

with st.form("input_surat"):
    col1, col2 = st.columns(2)
    with col1:
        nomor = st.text_input("Nomor Surat", placeholder="Contoh: 003/PENGKAB.PSH/II/2026")
        hal = st.text_input("Perihal", placeholder="Contoh: Edaran Agenda PSH")
    with col2:
        yth = st.text_input("Kepada Yth", placeholder="Contoh: Seluruh Warga PSH Tegal")
    
    arahan = st.text_area("Tulis Arahan/Inti Pesan:", placeholder="Contoh: Halal bihalal tgl 29 maret jam 10 pagi, tempat nyusul...")
    submit = st.form_submit_button("✨ Susun Surat Resmi")

if submit:
    with st.spinner("Sekretaris sedang merangkai kalimat..."):
        try:
            # PROMPT SUPER SEKRETARIS
            prompt = (f"Bertindaklah sebagai Sekretaris Organisasi Pencak Silat PSH Tegal. "
                      f"Buat isi surat resmi dari arahan ini: {arahan}. "
                      "ATURAN WAJIB: "
                      "1. Awali dengan paragraf pembuka yang luwes (Contoh: Sehubungan dengan agenda PSH Tegal, dengan ini kami sampaikan...). "
                      "2. Gunakan bahasa yang sopan dan persaudaraan. "
                      "3. Jika ada rincian (Waktu, Tempat, Acara), tulis dalam baris terpisah dengan format 'Label : Isi'. "
                      "4. Akhiri dengan paragraf penutup yang berisi harapan dan terima kasih (Contoh: Demikian surat ini kami buat, atas partisipasi sedulur semua kami haturkan Terimakasih). "
                      "5. JANGAN TULIS SALAM (Assalammualaikum) karena sudah ada di kertas surat. "
                      "6. Gunakan huruf kapital hanya di awal kata, jangan CAPSLOCK semua.")
            
            response = model.generate_content(prompt)
            st.session_state['draf_final'] = response.text.strip()
        except Exception as e:
            st.error(f"Gagal: {e}")

if 'draf_final' in st.session_state:
    st.subheader("📝 Review & Edit")
    isi_edit = st.text_area("Edit draf jika dirasa kurang pas:", value=st.session_state['draf_final'], height=350)
    st.session_state['draf_final'] = isi_edit

    if st.button("💾 Generate Word"):
        try:
            doc = Document("template_psh.docx")
            mapping = {"{{nomor}}": nomor, "{{hal}}": hal, "{{yth}}": yth, "{{tanggal}}": get_tanggal_indo()}
            
            for old, new in mapping.items():
                for p in doc.paragraphs:
                    if old in p.text:
                        for run in p.runs:
                            if old in run.text:
                                run.text = run.text.replace(old, str(new))
                                run.font.name = 'Times New Roman'

            format_surat_sekretaris(doc, "{{isi}}", st.session_state['draf_final'])
            
            output = io.BytesIO()
            doc.save(output)
            st.session_state['file_ok'] = output.getvalue()
            st.success("Surat berhasil dirakit!")
        except Exception as e:
            st.error(f"Error Word: {e}")

    if 'file_ok' in st.session_state:
        st.download_button("📥 Download Surat", data=st.session_state['file_ok'], file_name=f"Surat_PSH_{datetime.date.today()}.docx")
